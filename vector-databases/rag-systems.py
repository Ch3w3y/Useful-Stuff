#!/usr/bin/env python3
"""
Advanced RAG (Retrieval Augmented Generation) Systems with Vector Databases

Implementation of modern RAG systems for processing unstructured data with
vector embeddings, similarity search, and knowledge management capabilities.
"""

import asyncio
import json
import logging
import hashlib
import pickle
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field
from datetime import datetime
import uuid
import re
import os
from pathlib import Path

# Core libraries
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import PCA, TruncatedSVD

# Vector database libraries
try:
    import chromadb
    from chromadb.config import Settings
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False

try:
    import pinecone
    PINECONE_AVAILABLE = True
except ImportError:
    PINECONE_AVAILABLE = False

# Embedding models
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# Document processing
try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
    DOC_PROCESSING_AVAILABLE = True
except ImportError:
    DOC_PROCESSING_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document representation for RAG system"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    title: str = ""
    source: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[str] = field(default_factory=list)
    embeddings: Optional[np.ndarray] = None
    created_at: datetime = field(default_factory=datetime.now)

@dataclass
class QueryResult:
    """Result from similarity search"""
    document: Document
    chunk: str
    similarity_score: float
    chunk_index: int
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class RAGResponse:
    """Response from RAG system"""
    query: str
    answer: str
    sources: List[QueryResult]
    confidence: float
    reasoning: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

class DocumentProcessor:
    """Process various document types into text chunks"""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_text(self, text: str, source: str = "", title: str = "") -> Document:
        """Process plain text into document"""
        doc = Document(
            content=text,
            title=title or f"Text Document {datetime.now().strftime('%Y%m%d_%H%M%S')}",
            source=source,
            metadata={"type": "text", "length": len(text)}
        )
        doc.chunks = self.chunk_text(text)
        return doc
    
    def process_pdf(self, file_path: str) -> Document:
        """Process PDF file"""
        if not DOC_PROCESSING_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        doc = Document(
            content=text,
            title=Path(file_path).stem,
            source=file_path,
            metadata={
                "type": "pdf",
                "pages": len(pdf_reader.pages),
                "file_size": os.path.getsize(file_path)
            }
        )
        doc.chunks = self.chunk_text(text)
        return doc
    
    def process_docx(self, file_path: str) -> Document:
        """Process Word document"""
        if not DOC_PROCESSING_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        doc_obj = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc_obj.paragraphs])
        
        doc = Document(
            content=text,
            title=Path(file_path).stem,
            source=file_path,
            metadata={
                "type": "docx",
                "paragraphs": len(doc_obj.paragraphs),
                "file_size": os.path.getsize(file_path)
            }
        )
        doc.chunks = self.chunk_text(text)
        return doc
    
    def process_html(self, html_content: str, source: str = "") -> Document:
        """Process HTML content"""
        if not DOC_PROCESSING_AVAILABLE:
            raise ImportError("BeautifulSoup not available for HTML processing")
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.get_text() if title_tag else "HTML Document"
        
        # Extract text content
        text = soup.get_text()
        
        doc = Document(
            content=text,
            title=title,
            source=source,
            metadata={"type": "html", "length": len(text)}
        )
        doc.chunks = self.chunk_text(text)
        return doc
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        # Clean text
        text = re.sub(r'\s+', ' ', text.strip())
        
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + self.chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundaries
                    word_end = text.rfind(' ', start, end)
                    if word_end > start:
                        end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks

class EmbeddingManager:
    """Manage different embedding models and providers"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2", provider: str = "sentence-transformers"):
        self.model_name = model_name
        self.provider = provider
        self.model = None
        self.dimension = None
        self.setup_model()
    
    def setup_model(self):
        """Initialize embedding model"""
        if self.provider == "sentence-transformers" and SENTENCE_TRANSFORMERS_AVAILABLE:
            self.model = SentenceTransformer(self.model_name)
            self.dimension = self.model.get_sentence_embedding_dimension()
            logger.info(f"Loaded SentenceTransformer model: {self.model_name}")
            
        elif self.provider == "openai" and OPENAI_AVAILABLE:
            self.model = openai.OpenAI()
            self.dimension = 1536  # Default for text-embedding-ada-002
            logger.info("Loaded OpenAI embedding model")
            
        else:
            logger.warning(f"Embedding provider {self.provider} not available, using TF-IDF")
            self.model = TfidfVectorizer(max_features=512, stop_words='english')
            self.dimension = 512
            self.provider = "tfidf"
    
    def encode_text(self, texts: Union[str, List[str]]) -> np.ndarray:
        """Generate embeddings for text(s)"""
        if isinstance(texts, str):
            texts = [texts]
        
        if self.provider == "sentence-transformers":
            embeddings = self.model.encode(texts)
            return embeddings
            
        elif self.provider == "openai":
            embeddings = []
            for text in texts:
                response = self.model.embeddings.create(
                    model="text-embedding-ada-002",
                    input=text
                )
                embeddings.append(response.data[0].embedding)
            return np.array(embeddings)
            
        elif self.provider == "tfidf":
            if not hasattr(self.model, 'vocabulary_'):
                # Fit on first use
                self.model.fit(texts)
            embeddings = self.model.transform(texts).toarray()
            return embeddings
        
        else:
            raise ValueError(f"Unknown provider: {self.provider}")
    
    def compute_similarity(self, query_embedding: np.ndarray, 
                          document_embeddings: np.ndarray) -> np.ndarray:
        """Compute cosine similarity between query and documents"""
        return cosine_similarity([query_embedding], document_embeddings)[0]

class VectorDatabase:
    """Abstract base class for vector databases"""
    
    def __init__(self, collection_name: str = "documents"):
        self.collection_name = collection_name
    
    def add_documents(self, documents: List[Document], embeddings: np.ndarray):
        """Add documents with embeddings"""
        raise NotImplementedError
    
    def search(self, query_embedding: np.ndarray, k: int = 5) -> List[Dict]:
        """Search for similar documents"""
        raise NotImplementedError
    
    def delete_document(self, document_id: str):
        """Delete document by ID"""
        raise NotImplementedError

class ChromaVectorDB(VectorDatabase):
    """ChromaDB implementation"""
    
    def __init__(self, collection_name: str = "documents", persist_directory: str = "./chroma_db"):
        super().__init__(collection_name)
        
        if not CHROMA_AVAILABLE:
            raise ImportError("ChromaDB not available")
        
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection = self.client.get_or_create_collection(name=collection_name)
        
    def add_documents(self, documents: List[Document], embeddings: np.ndarray):
        """Add documents to ChromaDB"""
        ids = []
        metadatas = []
        docs = []
        embeds = []
        
        embed_index = 0
        for doc in documents:
            for i, chunk in enumerate(doc.chunks):
                chunk_id = f"{doc.id}_{i}"
                ids.append(chunk_id)
                
                metadata = {
                    "document_id": doc.id,
                    "title": doc.title,
                    "source": doc.source,
                    "chunk_index": i,
                    "created_at": doc.created_at.isoformat(),
                    **doc.metadata
                }
                metadatas.append(metadata)
                docs.append(chunk)
                embeds.append(embeddings[embed_index].tolist())
                embed_index += 1
        
        self.collection.add(
            ids=ids,
            embeddings=embeds,
            metadatas=metadatas,
            documents=docs
        )
    
    def search(self, query_embedding: np.ndarray, k: int = 5) -> List[Dict]:
        """Search ChromaDB"""
        results = self.collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=k
        )
        
        search_results = []
        for i in range(len(results['ids'][0])):
            search_results.append({
                'id': results['ids'][0][i],
                'document': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'distance': results['distances'][0][i] if 'distances' in results else None
            })
        
        return search_results

class FAISSVectorDB(VectorDatabase):
    """FAISS implementation"""
    
    def __init__(self, collection_name: str = "documents", dimension: int = 384):
        super().__init__(collection_name)
        
        if not FAISS_AVAILABLE:
            raise ImportError("FAISS not available")
        
        self.dimension = dimension
        self.index = faiss.IndexFlatIP(dimension)  # Inner product (cosine similarity)
        self.documents = {}  # Store document metadata
        self.document_mapping = []  # Map index to document info
    
    def add_documents(self, documents: List[Document], embeddings: np.ndarray):
        """Add documents to FAISS index"""
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        
        # Add to index
        start_idx = self.index.ntotal
        self.index.add(embeddings)
        
        # Store metadata
        embed_index = 0
        for doc in documents:
            self.documents[doc.id] = doc
            for i, chunk in enumerate(doc.chunks):
                self.document_mapping.append({
                    'document_id': doc.id,
                    'chunk_index': i,
                    'chunk': chunk
                })
                embed_index += 1
    
    def search(self, query_embedding: np.ndarray, k: int = 5) -> List[Dict]:
        """Search FAISS index"""
        # Normalize query embedding
        query_embedding = query_embedding.reshape(1, -1).astype(np.float32)
        faiss.normalize_L2(query_embedding)
        
        # Search
        scores, indices = self.index.search(query_embedding, k)
        
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(self.document_mapping):
                doc_info = self.document_mapping[idx]
                doc = self.documents[doc_info['document_id']]
                
                results.append({
                    'id': f"{doc_info['document_id']}_{doc_info['chunk_index']}",
                    'document': doc_info['chunk'],
                    'metadata': {
                        'document_id': doc.id,
                        'title': doc.title,
                        'source': doc.source,
                        'chunk_index': doc_info['chunk_index'],
                        **doc.metadata
                    },
                    'score': float(scores[0][i])
                })
        
        return results

class RAGSystem:
    """Complete RAG (Retrieval Augmented Generation) system"""
    
    def __init__(self, vector_db: VectorDatabase, embedding_manager: EmbeddingManager,
                 llm_provider: str = "openai", model_name: str = "gpt-3.5-turbo"):
        self.vector_db = vector_db
        self.embedding_manager = embedding_manager
        self.llm_provider = llm_provider
        self.model_name = model_name
        self.document_processor = DocumentProcessor()
        self.setup_llm()
    
    def setup_llm(self):
        """Setup language model for generation"""
        if self.llm_provider == "openai" and OPENAI_AVAILABLE:
            self.llm_client = openai.OpenAI()
        else:
            logger.warning(f"LLM provider {self.llm_provider} not available")
            self.llm_client = None
    
    def add_document(self, document: Document):
        """Add single document to RAG system"""
        return self.add_documents([document])
    
    def add_documents(self, documents: List[Document]):
        """Add multiple documents to RAG system"""
        # Generate embeddings for all chunks
        all_chunks = []
        for doc in documents:
            all_chunks.extend(doc.chunks)
        
        if all_chunks:
            embeddings = self.embedding_manager.encode_text(all_chunks)
            self.vector_db.add_documents(documents, embeddings)
            logger.info(f"Added {len(documents)} documents with {len(all_chunks)} chunks")
    
    def add_text_document(self, text: str, title: str = "", source: str = "") -> Document:
        """Add text document directly"""
        doc = self.document_processor.process_text(text, source, title)
        self.add_document(doc)
        return doc
    
    def add_file(self, file_path: str) -> Document:
        """Add document from file"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            doc = self.document_processor.process_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            doc = self.document_processor.process_docx(str(file_path))
        elif file_path.suffix.lower() in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            doc = self.document_processor.process_text(text, str(file_path), file_path.stem)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        self.add_document(doc)
        return doc
    
    def retrieve_contexts(self, query: str, k: int = 5) -> List[QueryResult]:
        """Retrieve relevant contexts for query"""
        # Generate query embedding
        query_embedding = self.embedding_manager.encode_text([query])[0]
        
        # Search vector database
        search_results = self.vector_db.search(query_embedding, k)
        
        # Convert to QueryResult objects
        contexts = []
        for result in search_results:
            # Create dummy document (in real system, retrieve from storage)
            doc = Document(
                id=result['metadata']['document_id'],
                title=result['metadata']['title'],
                source=result['metadata']['source']
            )
            
            query_result = QueryResult(
                document=doc,
                chunk=result['document'],
                similarity_score=result.get('score', 1.0 - result.get('distance', 0.0)),
                chunk_index=result['metadata']['chunk_index'],
                metadata=result['metadata']
            )
            contexts.append(query_result)
        
        return contexts
    
    def generate_answer(self, query: str, contexts: List[QueryResult]) -> str:
        """Generate answer using LLM with retrieved contexts"""
        if not self.llm_client:
            # Fallback: simple context-based response
            context_text = "\n\n".join([ctx.chunk for ctx in contexts[:3]])
            return f"Based on the available information:\n\n{context_text}\n\nThis information is relevant to your query: {query}"
        
        # Prepare context for LLM
        context_text = ""
        for i, ctx in enumerate(contexts[:5]):
            context_text += f"Context {i+1} (from {ctx.document.title}):\n{ctx.chunk}\n\n"
        
        # Create prompt
        prompt = f"""Based on the following context information, please provide a comprehensive answer to the user's question.

Context Information:
{context_text}

User Question: {query}

Please provide a detailed answer based on the context provided. If the context doesn't contain enough information to fully answer the question, please indicate what information is missing. Cite the relevant context sources when appropriate.

Answer:"""
        
        try:
            response = self.llm_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that answers questions based on provided context."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            return "I apologize, but I encountered an error while generating the answer. Please try again."
    
    def query(self, question: str, k: int = 5) -> RAGResponse:
        """Complete RAG query pipeline"""
        logger.info(f"Processing query: {question}")
        
        # Retrieve relevant contexts
        contexts = self.retrieve_contexts(question, k)
        
        # Generate answer
        answer = self.generate_answer(question, contexts)
        
        # Calculate confidence based on similarity scores
        if contexts:
            avg_similarity = np.mean([ctx.similarity_score for ctx in contexts])
            confidence = float(avg_similarity)
        else:
            confidence = 0.0
        
        response = RAGResponse(
            query=question,
            answer=answer,
            sources=contexts,
            confidence=confidence,
            reasoning=f"Retrieved {len(contexts)} relevant contexts with average similarity {confidence:.3f}",
            metadata={
                "timestamp": datetime.now().isoformat(),
                "model": self.model_name,
                "embedding_model": self.embedding_manager.model_name
            }
        )
        
        return response
    
    def batch_query(self, questions: List[str], k: int = 5) -> List[RAGResponse]:
        """Process multiple queries"""
        responses = []
        for question in questions:
            response = self.query(question, k)
            responses.append(response)
        return responses

class KnowledgeManager:
    """Advanced knowledge management with multiple RAG systems"""
    
    def __init__(self):
        self.rag_systems: Dict[str, RAGSystem] = {}
        self.document_index: Dict[str, List[str]] = {}  # Document ID to RAG system mapping
    
    def create_knowledge_base(self, name: str, vector_db_type: str = "faiss",
                            embedding_model: str = "all-MiniLM-L6-v2") -> RAGSystem:
        """Create a new knowledge base"""
        # Setup embedding manager
        embedding_manager = EmbeddingManager(model_name=embedding_model)
        
        # Setup vector database
        if vector_db_type == "chroma" and CHROMA_AVAILABLE:
            vector_db = ChromaVectorDB(collection_name=name)
        elif vector_db_type == "faiss" and FAISS_AVAILABLE:
            vector_db = FAISSVectorDB(collection_name=name, dimension=embedding_manager.dimension)
        else:
            logger.warning(f"Vector DB {vector_db_type} not available, using in-memory fallback")
            vector_db = InMemoryVectorDB(collection_name=name)
        
        # Create RAG system
        rag_system = RAGSystem(vector_db, embedding_manager)
        self.rag_systems[name] = rag_system
        
        logger.info(f"Created knowledge base: {name}")
        return rag_system
    
    def get_knowledge_base(self, name: str) -> Optional[RAGSystem]:
        """Get existing knowledge base"""
        return self.rag_systems.get(name)
    
    def add_documents_to_kb(self, kb_name: str, documents: List[Document]):
        """Add documents to specific knowledge base"""
        if kb_name not in self.rag_systems:
            raise ValueError(f"Knowledge base {kb_name} not found")
        
        rag_system = self.rag_systems[kb_name]
        rag_system.add_documents(documents)
        
        # Update document index
        for doc in documents:
            if doc.id not in self.document_index:
                self.document_index[doc.id] = []
            self.document_index[doc.id].append(kb_name)
    
    def cross_kb_query(self, question: str, kb_names: Optional[List[str]] = None, k: int = 3) -> Dict[str, RAGResponse]:
        """Query across multiple knowledge bases"""
        if kb_names is None:
            kb_names = list(self.rag_systems.keys())
        
        responses = {}
        for kb_name in kb_names:
            if kb_name in self.rag_systems:
                response = self.rag_systems[kb_name].query(question, k)
                responses[kb_name] = response
        
        return responses
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get knowledge management statistics"""
        stats = {
            "knowledge_bases": len(self.rag_systems),
            "total_documents": len(self.document_index),
            "kb_details": {}
        }
        
        for kb_name, rag_system in self.rag_systems.items():
            # This would need to be implemented in each vector DB
            stats["kb_details"][kb_name] = {
                "embedding_model": rag_system.embedding_manager.model_name,
                "vector_db_type": type(rag_system.vector_db).__name__
            }
        
        return stats

class InMemoryVectorDB(VectorDatabase):
    """Simple in-memory vector database for testing"""
    
    def __init__(self, collection_name: str = "documents"):
        super().__init__(collection_name)
        self.embeddings = []
        self.documents = []
        self.metadata = []
    
    def add_documents(self, documents: List[Document], embeddings: np.ndarray):
        """Add documents to memory"""
        embed_index = 0
        for doc in documents:
            for i, chunk in enumerate(doc.chunks):
                self.embeddings.append(embeddings[embed_index])
                self.documents.append(chunk)
                self.metadata.append({
                    'document_id': doc.id,
                    'title': doc.title,
                    'source': doc.source,
                    'chunk_index': i,
                    **doc.metadata
                })
                embed_index += 1
        
        self.embeddings = np.array(self.embeddings) if self.embeddings else np.array([])
    
    def search(self, query_embedding: np.ndarray, k: int = 5) -> List[Dict]:
        """Search in memory"""
        if len(self.embeddings) == 0:
            return []
        
        # Compute similarities
        similarities = cosine_similarity([query_embedding], self.embeddings)[0]
        
        # Get top k results
        top_indices = np.argsort(similarities)[::-1][:k]
        
        results = []
        for idx in top_indices:
            results.append({
                'id': f"{self.metadata[idx]['document_id']}_{self.metadata[idx]['chunk_index']}",
                'document': self.documents[idx],
                'metadata': self.metadata[idx],
                'score': float(similarities[idx])
            })
        
        return results

# Demo and example usage
def demo_rag_system():
    """Demonstrate RAG system capabilities"""
    print("=== RAG System Demo ===\n")
    
    # Create knowledge manager
    km = KnowledgeManager()
    
    # Create knowledge base
    print("1. Creating knowledge base...")
    kb = km.create_knowledge_base("demo_kb", vector_db_type="memory")  # Use in-memory for demo
    
    # Add sample documents
    print("2. Adding sample documents...")
    
    sample_docs = [
        "Artificial Intelligence (AI) is intelligence demonstrated by machines, in contrast to the natural intelligence displayed by humans and animals. AI research has been highly successful in developing effective techniques for solving problems, from game playing to medical diagnosis.",
        
        "Machine Learning is a subset of artificial intelligence that focuses on algorithms that allow computers to learn and improve from experience without being explicitly programmed. Common types include supervised learning, unsupervised learning, and reinforcement learning.",
        
        "Deep Learning is a machine learning technique that teaches computers to learn by example, the way humans do. It uses neural networks with multiple layers to model and understand complex patterns in data. Deep learning is behind many recent AI breakthroughs.",
        
        "Data Science is an interdisciplinary field that uses scientific methods, processes, algorithms and systems to extract knowledge and insights from structured and unstructured data. It combines statistics, machine learning, and domain expertise.",
        
        "Natural Language Processing (NLP) is a branch of AI that helps computers understand, interpret and manipulate human language. NLP draws from many disciplines, including computer science and computational linguistics."
    ]
    
    for i, text in enumerate(sample_docs):
        title = f"Document {i+1}: {text.split('.')[0][:50]}..."
        kb.add_text_document(text, title=title, source=f"sample_doc_{i+1}")
    
    # Test queries
    print("3. Testing queries...")
    
    queries = [
        "What is artificial intelligence?",
        "How does machine learning work?",
        "What are the types of machine learning?",
        "What is the relationship between AI and deep learning?",
        "How is data science different from machine learning?"
    ]
    
    for query in queries:
        print(f"\nQuery: {query}")
        response = kb.query(query, k=3)
        print(f"Answer: {response.answer[:200]}...")
        print(f"Confidence: {response.confidence:.3f}")
        print(f"Sources: {len(response.sources)} documents")
        
        if response.sources:
            print("Top source:", response.sources[0].document.title[:50])
    
    # Cross-KB demonstration
    print("\n4. Creating second knowledge base...")
    kb2 = km.create_knowledge_base("tech_kb")
    
    tech_docs = [
        "Cloud Computing is the delivery of computing services over the internet. It includes servers, storage, databases, networking, software, analytics, and intelligence.",
        
        "Blockchain is a distributed ledger technology that maintains a continuously growing list of records, called blocks, which are linked and secured using cryptography.",
        
        "Internet of Things (IoT) refers to the network of physical objects embedded with sensors, software, and connectivity to collect and exchange data."
    ]
    
    for i, text in enumerate(tech_docs):
        title = f"Tech Doc {i+1}: {text.split('.')[0][:50]}..."
        kb2.add_text_document(text, title=title, source=f"tech_doc_{i+1}")
    
    print("5. Cross-knowledge base query...")
    cross_responses = km.cross_kb_query("What are emerging technologies?")
    
    for kb_name, response in cross_responses.items():
        print(f"\nKnowledge Base: {kb_name}")
        print(f"Answer: {response.answer[:150]}...")
        print(f"Confidence: {response.confidence:.3f}")
    
    # Statistics
    print("\n6. Knowledge management statistics:")
    stats = km.get_statistics()
    print(json.dumps(stats, indent=2))
    
    return km

if __name__ == "__main__":
    # Run demo
    demo_rag_system() 